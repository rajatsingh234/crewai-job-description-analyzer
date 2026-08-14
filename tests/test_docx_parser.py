from docx import Document
from crewai_job_description_analyzer.document_processing.docx_parser import (
    extract_text_from_docx,
)


def test_docx_parser():
    text = extract_text_from_docx("tests/Rajat_Singh_Offer_letter.docx")

    print("\nEXTRACTED RESUME TEXT:\n")
    print(text)

    assert text.strip()

def test_docx_parser_extracts_table(tmp_path):
    document = Document()

    document.add_paragraph("Resume")

    table = document.add_table(rows=2, cols=2)

    table.cell(0, 0).text = "Skills"
    table.cell(0, 1).text = "Python"

    table.cell(1, 0).text = "Framework"
    table.cell(1, 1).text = "Pytest"

    file_path = tmp_path / "table_resume.docx"
    document.save(file_path)

    text = extract_text_from_docx(str(file_path))

    assert "Skills | Python" in text
    assert "Framework | Pytest" in text