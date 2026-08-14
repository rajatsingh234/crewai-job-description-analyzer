from docx import Document


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX resume.

    Extracts text from:
    - Paragraphs
    - Tables

    Args:
        file_path: Path to the DOCX file.

    Returns:
        Extracted text from the document.
    """

    document = Document(file_path)

    sections: list[str] = []

    # Extract paragraphs
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            sections.append(text)

    # Extract tables
    for table in document.tables:
        for row in table.rows:
            cells = []

            for cell in row.cells:
                text = cell.text.strip()

                if text:
                    cells.append(text)

            if cells:
                sections.append(" | ".join(cells))

    return "\n\n".join(sections)